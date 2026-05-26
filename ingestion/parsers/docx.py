"""UrivDocs — ingestion/parsers/docx.py"""

from pathlib import Path
from typing import List

from docx import Document
from loguru import logger

from ingestion.parsers.base import BasePage


class DocxParser:
    def parse(self, path: Path) -> List[BasePage]:
        doc = Document(str(path))
        pages: List[BasePage] = []

        current_section = None
        buffer: List[str] = []
        page_num = 1

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""

            if style.startswith("Heading"):
                if buffer:
                    pages.append(BasePage(
                        text="\n".join(buffer),
                        page_number=page_num,
                        section=current_section,
                    ))
                    buffer = []
                    page_num += 1

                current_section = para.text.strip()
            elif para.text.strip():
                buffer.append(para.text.strip())

        if buffer:
            pages.append(BasePage(
                text="\n".join(buffer),
                page_number=page_num,
                section=current_section,
            ))

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                pages.append(BasePage(
                    text="\n".join(rows),
                    page_number=page_num + 1,
                    section="Table",
                ))

        logger.info(f"Extracted {len(pages)} sections from DOCX")
        return pages
